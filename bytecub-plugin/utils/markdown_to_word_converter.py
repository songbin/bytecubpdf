import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
import markdown
from pylatexenc.latex2text import LatexNodes2Text
from pylatexenc.latexwalker import LatexWalker
import matplotlib.pyplot as plt
from io import BytesIO
from PIL import Image


class MarkdownToWordConverter:
    def __init__(self, md_file_path, docx_file_path):
        self.md_file_path = md_file_path
        self.docx_file_path = docx_file_path
        self.doc = Document()

    def convert(self):
        """主转换函数"""
        # 读取 Markdown 文件
        with open(self.md_file_path, 'r', encoding='utf-8') as f:
            md_text = f.read()

        # 转换为 HTML
        html = markdown.markdown(md_text, extensions=['tables'], output_format='html5')
        soup = BeautifulSoup(html, 'html.parser')

        # 解析并转换 HTML 元素
        for element in soup.find_all(recursive=False):
            self._parse_element(element)

        # 保存为 Word 文档
        docx_dir = os.path.dirname(self.docx_file_path)
        if docx_dir and not os.path.exists(docx_dir):
            os.makedirs(docx_dir)
        self.doc.save(self.docx_file_path)
        print(f"✅ 转换完成：{self.docx_file_path}")

    def _parse_element(self, element):
        """解析 HTML 元素并添加到 Word 文档"""
        if element.name == 'h1':
            self._add_heading(element.get_text(strip=True), level=1)
        elif element.name == 'h2':
            self._add_heading(element.get_text(strip=True), level=2)
        elif element.name == 'h3':
            self._add_heading(element.get_text(strip=True), level=3)
        elif element.name == 'p':
            self._add_paragraph(element)
        elif element.name == 'ul':
            self._add_list(element, ordered=False)
        elif element.name == 'ol':
            self._add_list(element, ordered=True)
        elif element.name == 'table':
            self._add_table(element)
        elif element.name == 'a':
            self._add_hyperlink(element)
        elif element.name == 'img':
            self._add_image(element)
        elif element.name == 'pre':
            self._add_code_block(element)
        elif element.name == 'blockquote':
            self._add_blockquote(element)
        elif element.name is not None:
            self._add_paragraph(element)

    def _add_heading(self, text, level):
        self.doc.add_heading(text, level=level)

    def _add_paragraph(self, element):
        p = self.doc.add_paragraph()
        for child in element.children:
            if child.name == 'span' and 'math' in child.get('class', []):
                latex = child.get_text(strip=True)
                self._add_latex_equation(p, latex)
            elif child.name == 'a':
                self._add_hyperlink_to_paragraph(p, child)
            elif child.name == 'strong':
                run = p.add_run(child.get_text(strip=True))
                run.bold = True
            elif child.name == 'em':
                run = p.add_run(child.get_text(strip=True))
                run.italic = True
            elif child.name == 'br':
                p.add_run('\n')
            elif child.string:
                p.add_run(child.string)

    def _add_list(self, element, ordered):
        items = [li.get_text(strip=True) for li in element.find_all('li')]
        style = 'ListNumber' if ordered else 'ListBullet'
        for item in items:
            p = self.doc.add_paragraph(item, style=style)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_table(self, table_element, parent_cell=None):
        rows = table_element.find_all('tr')
        if not rows:
            return

        max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
        
        # 根据是否为嵌套表格选择不同的添加方式
        if parent_cell is None:
            table = self.doc.add_table(rows=len(rows), cols=max_cols)
        else:
            table = parent_cell.add_table(rows=len(rows), cols=max_cols)
        
        table.style = 'Table Grid'
        self._set_table_style(table)

        cell_map = {}
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            col_idx = 0
            for j, cell in enumerate(cells):
                while (i, col_idx) in cell_map:
                    col_idx += 1

                text = cell.get_text(strip=True)
                rowspan = int(cell.get('rowspan', 1))
                colspan = int(cell.get('colspan', 1))

                merged_cell = table.cell(i, col_idx)
                if rowspan > 1 or colspan > 1:
                    end_row = i + rowspan - 1
                    end_col = col_idx + colspan - 1
                    merged_cell = table.cell(i, col_idx).merge(table.cell(end_row, end_col))

                sub_tables = cell.find_all('table')
                if sub_tables:
                    for sub_table in sub_tables:
                        self._add_table(sub_table, parent_cell=merged_cell)
                else:
                    merged_cell.text = text

                for r in range(i, i + rowspan):
                    for c in range(col_idx, col_idx + colspan):
                        cell_map[(r, c)] = True

                col_idx += colspan
    def _set_table_style(self, table, border_color="000000", bg_color="F0F0F0", align="center"):
        """设置表格样式（修正后的版本）"""
        tbl = table._tbl  # 获取底层XML元素

        # 设置边框
        borders = tbl.xpath('.//w:tblBorders')
        if not borders:
            borders = OxmlElement('w:tblBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:color'), border_color)
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')
                borders.append(b)
            tbl.tblPr.append(borders)

        # 设置表格背景颜色（表级）
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), bg_color)
        # 查找或创建 tblPr
        tbl_pr = tbl.xpath('.//w:tblPr')
        if not tbl_pr:
            tbl_pr = OxmlElement('w:tblPr')
            tbl.insert(0, tbl_pr)  # 插入到表格的最开始
        tbl_pr = tbl.tblPr  # 重新获取
        tbl_pr.append(shading)

        # 设置对齐方式
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), align)
        tbl_pr.append(jc)    
    def _add_hyperlink(self, element):
        href = element.get('href')
        text = element.get_text(strip=True)
        if href and text:
            p = self.doc.add_paragraph()
            self._add_hyperlink_to_paragraph(p, element)

    def _add_hyperlink_to_paragraph(self, paragraph, element):
        href = element.get('href')
        text = element.get_text(strip=True)
        if href and text:
            part = paragraph.part
            r_id = part.relate_to(href, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)

            new_run = OxmlElement('w:r')
            new_run.text = text
            new_run.set(qn('w:b'), 'true')
            new_run.set(qn('w:color'), '0000FF')

            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)

    def _add_image(self, element):
        src = element.get('src')
        if src and os.path.exists(src):
            self.doc.add_picture(src, width=Inches(5.5))

    def _add_code_block(self, element):
        code = element.get_text(strip=True)
        p = self.doc.add_paragraph(code)
        p.style = 'No Spacing'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    def _add_blockquote(self, element):
        text = element.get_text(strip=True)
        p = self.doc.add_paragraph(f"“{text}”")
        p.style = 'Intense Quote'

    def _add_latex_equation(self, paragraph, latex_str):
        """将 LaTeX 公式插入 Word 段落中"""
        try:
            walker = LatexWalker(latex_str)
            nodes = walker.parse()
            text = LatexNodes2Text().nodelist_to_text(nodes)

            math = OxmlElement('m:oMath')
            math.set(qn('m:val'), text)
            paragraph._p.append(math)
        except Exception:
            # 回退为插入图片公式
            plt.figure(figsize=(6, 1))
            plt.text(0.5, 0.5, f"${latex_str}$", fontsize=16, ha='center')
            plt.axis('off')
            img_data = BytesIO()
            plt.savefig(img_data, format='png', bbox_inches='tight', dpi=300)
            plt.close()

            img = Image.open(img_data)
            img = img.convert("RGBA")
            img_data = BytesIO()
            img.save(img_data, format='PNG')

            p = paragraph
            p.add_run().add_picture(img_data, width=Inches(5.5))